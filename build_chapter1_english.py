from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE = Path(r"D:\组会\2026\8.15\新建 DOCX 文档.docx")
OUTPUT = Path(r"D:\组会\2026\8.15\Chapter_1_English_Polished.docx")


BODY_PARAGRAPHS = [
    "The evolution of sixth-generation (6G) mobile communication systems is transforming wireless networks from conventional communication infrastructure into platforms for integrated sensing and communication (ISAC) [Reference]. Beyond carrying information, radio signals are perturbed by human motion, object displacement, and changes in the surrounding environment. The communication link can therefore serve as a sensing medium. This capability has stimulated extensive research on radio-based human activity recognition, gesture recognition, and person identification [Reference]. Compared with cameras and wearable sensors, wireless sensing operates independently of illumination, requires no dedicated device to be carried by the user, and reduces the collection of visually identifiable information. These properties make it attractive for smart homes, health monitoring, and human-computer interaction [Reference].",
    "Among the available wireless modalities, channel state information (CSI) has become a principal signal source for sensing. CSI characterizes the wireless channel response across subcarriers and antenna links; variations in its amplitude and phase encode perturbations to multipath propagation caused by human motion. Early systems extracted amplitude statistics, Doppler shifts, and time-frequency descriptors from CSI or received signals and classified activities and gestures using conventional models, as exemplified by E-eyes, CARM, and Widar [Reference]. More recent work applies convolutional neural networks (CNNs), recurrent neural networks (RNNs), Transformers, and other sequence models to learn spatiotemporal representations directly from CSI [Reference]. Systems such as Widar3.0 and SignFi, together with benchmarking frameworks such as SenseFi, have expanded cross-domain gesture recognition, human activity recognition, and standardized model evaluation [Reference]. These advances show that learning-based CSI sensing can support complex human-sensing tasks. They also expose a fundamental limitation of model-centric evaluation: recognition performance depends on the signal representation presented to the model, not on the model alone.",
    "Raw CSI is rarely a model-ready input. Measurements contain task-relevant channel variations induced by human motion, but they also reflect transceiver nonidealities, clock asynchrony, environmental interference, and other nuisance components. Hardware asynchrony can introduce phase offsets and subcarrier-dependent phase trends; impulsive interference can produce anomalous amplitude samples; and acquisition platforms can differ in both subcarrier count and sampling rate. CSI therefore typically undergoes a sequence of preprocessing operations before model training. These operations are intended to suppress task-irrelevant disturbances and produce a numerically and structurally compatible input representation.",
    "Wireless sensing systems employ a broad range of such operations, including denoising, outlier detection and correction, phase calibration to mitigate hardware-induced distortion, normalization to control signal scale, and interpolation to harmonize the subcarrier dimension [Reference]. None of these operations is intrinsically benign. A denoiser can suppress high-frequency noise while attenuating rapid motion-induced variations. Phase calibration can remove hardware drift but alter motion-related phase structure. The window and threshold used for outlier correction can determine whether a transient is treated as interference or as informative motion. Preprocessing thus defines which components of CSI are preserved, suppressed, or transformed, and consequently constrains the features that a downstream model can learn.",
    "Nevertheless, wireless sensing research remains predominantly organized around model architecture and end-task recognition performance [Reference]. New convolutional, recurrent, attention-based, and sequence-modeling approaches are usually compared using aggregate metrics such as accuracy and F1 score. The preprocessing pipeline, by contrast, is often treated as a fixed implementation detail, even though studies differ in their denoising, phase processing, normalization, and related choices. Performance differences between two sensing systems may therefore reflect not only the learning models but also the preprocessing strategies and their interactions with those models. The original Sensing Data Protocol (SDP) study showed that inconsistent data representations and processing pipelines can confound algorithmic gains with implementation differences, thereby undermining fair comparison and reproducibility [Reference].",
    "This confounding leaves three questions unresolved. First, to what extent do commonly used CSI preprocessing operations improve end-task performance across heterogeneous datasets? A denoising, phase-calibration, or normalization method that is effective for one dataset may not retain the same benefit across hardware platforms, sensing tasks, sampling regimes, and CSI representations. Second, how strongly does the effect depend on the downstream model? Models differ in feature-extraction capacity and inductive bias: some may learn to attenuate particular input disturbances, whereas others may rely heavily on a carefully processed representation. Evidence obtained with a single model cannot therefore characterize an entire sensing pipeline. Third, how do successive preprocessing components interact? Operations may be complementary, redundant, or destructive when composed, and an apparent gain from one configuration does not establish that any constituent operation is beneficial in isolation. The central problem is consequently not to identify a universally optimal pipeline, but to determine when an operation helps, when its benefit is limited, and when it degrades sensing information.",
    "SDP provides an appropriate foundation for answering these questions. It standardizes heterogeneous wireless sensing data at the protocol level through common data interfaces, CSI representations, and training and evaluation procedures, thereby reducing confounding from hardware and implementation differences. Its primary objective, however, is to establish a unified, stable, and reproducible protocol rather than to compare the choices and operating conditions of alternative preprocessing algorithms and models [Reference]. SDP therefore addresses how wireless sensing methods can be evaluated through a common interface, but it leaves a more focused question open: how should CSI preprocessing be selected within a standardized protocol and evaluation environment? We extend the preprocessing and model layers of SDP so that preprocessing operations and learning models become explicit, configurable, and comparable experimental factors.",
    "We evaluate five principal preprocessing stages: denoising, outlier handling, phase calibration, normalization, and interpolation. The experiments cover four CSI datasets with distinct tasks and data properties: Widar, Gait, XRF55, and ElderAL. The evaluation follows a two-stage design. Stage I applies six representative SDP preprocessing presets to the complete model set, allowing us to quantify how model performance changes both within and across preprocessing conditions and to identify model-preprocessing dependence. Rather than selecting the single highest-scoring model from one run, this stage identifies a competitive and representative model for each dataset. Stage II fixes that model and evaluates fine-grained combinations of the five preprocessing stages under identical data splits and training settings, followed by targeted analyses of poorly performing configurations. This design limits two common inference errors: attributing all performance differences to a single model and generalizing the best configuration from one dataset into a universal preprocessing rule.",
    "The experiments reveal three consistent phenomena. First, preprocessing can materially alter end-task performance, but the effect is not uniformly positive; commonly used operations can cause pronounced degradation under particular data and task conditions. Second, effectiveness depends on whether an operation is compatible with the physical representation, sampling characteristics, and temporal dynamics of the CSI. Motion information is weakened when task-relevant variations are misclassified as noise, outliers, or hardware trends. Third, preprocessing components interact: a downstream operation can amplify, mask, or compensate for information changes introduced upstream. More preprocessing is therefore not necessarily better, and no fixed pipeline is appropriate for every sensing scenario. These findings motivate a data- and task-aware set of principles for selecting CSI preprocessing operations.",
]


CONTRIBUTIONS = [
    ("A preprocessing-centric extension of SDP.",
     "We develop an evaluation framework in which denoising, outlier handling, phase calibration, normalization, and interpolation are configurable and composable stages. In contrast to benchmarks centered on model architecture or a fixed preprocessing pipeline, the framework makes preprocessing itself a controlled research object within a common experimental interface."),
    ("A systematic cross-dataset and cross-model empirical evaluation.",
     "We compare multiple learning models and representative preprocessing pipelines on four public datasets with different sensing tasks and data properties: Widar, Gait, XRF55, and ElderAL. We then evaluate fine-grained combinations of the five preprocessing stages. The resulting evidence characterizes how preprocessing effects vary across datasets and models under controlled evaluation conditions."),
    ("Mechanistic analysis of preprocessing gains and failures.",
     "Beyond ranking configurations by end-task performance, we combine signal-characteristic analysis with targeted ablations to examine why individual operations improve or degrade recognition. The analysis tests, for example, whether denoising attenuates rapid motion-induced variations, whether phase calibration damages motion-related phase structure, and whether successive stages amplify earlier information loss."),
    ("Practical principles for selecting CSI preprocessing operations.",
     "We synthesize the empirical and mechanistic findings into selection principles that account for CSI representation, sampling characteristics, task dynamics, and model dependence. The resulting guidance supports defensible pipeline configuration and reproducible experimentation without presuming a single universally optimal preprocessing pipeline."),
]


FIGURE_CAPTION = (
    "Fig. 1. Motivation for preprocessing-centric evaluation of WiFi CSI sensing. "
    "(a) Raw CSI contains motion-related variations, hardware-induced drift, and noise. "
    "(b) Preprocessing operators encode different assumptions about these components. "
    "(c) An operator may preserve task-relevant structure or remove it as interference. "
    "(d) Consequently, the same learning model can support different empirical conclusions under different preprocessing pipelines."
)


def set_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def configure_style(style, size, bold=False, italic=False, color=(0, 0, 0),
                    before=0, after=0, line=1.0, keep_next=False):
    set_font(style, size=size, bold=bold, italic=italic, color=color)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.widow_control = True


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def add_numbering_style(doc, style, num_id=9):
    ppr = style.element.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.insert(0, num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def build_document():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    source_doc = Document(SOURCE)
    source_image = None
    image_para = source_doc.paragraphs[27]
    for blip in image_para._element.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed"))
        source_image = source_doc.part.rels[rid].target_part.blob
        break

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    configure_style(styles["Normal"], 11, after=6, line=1.15)
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    configure_style(styles["Title"], 16, bold=True, after=12, line=1.05, keep_next=True)
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_style(styles["Subtitle"], 10.5, italic=True, color=(80, 80, 80), after=16, line=1.05)
    styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_style(styles["Heading 1"], 14, bold=True, before=6, after=10, line=1.05, keep_next=True)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    configure_style(styles["Heading 2"], 12, bold=True, before=10, after=5, line=1.05, keep_next=True)
    configure_style(styles["Caption"], 9, italic=False, after=8, line=1.0, keep_next=False)
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Contribution" not in [s.name for s in styles]:
        contribution_style = styles.add_style("Contribution", WD_STYLE_TYPE.PARAGRAPH)
    else:
        contribution_style = styles["Contribution"]
    configure_style(contribution_style, 10.5, after=5, line=1.1)
    contribution_style.paragraph_format.left_indent = Cm(0.75)
    contribution_style.paragraph_format.first_line_indent = Cm(-0.75)
    contribution_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_numbering_style(doc, contribution_style)

    title = doc.add_paragraph(style="Title")
    set_font(title.add_run("Systematic Evaluation of Pluggable Preprocessing Pipelines for Heterogeneous WiFi CSI Sensing Datasets"), size=16, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_font(subtitle.add_run("Polished English Translation of Chapter 1"), size=10.5, italic=True, color=(80, 80, 80))

    h1 = doc.add_paragraph(style="Heading 1")
    set_font(h1.add_run("1. Introduction"), size=14, bold=True)

    for text in BODY_PARAGRAPHS[:8]:
        p = doc.add_paragraph(style="Normal")
        set_font(p.add_run(text), size=11)

    h2 = doc.add_paragraph(style="Heading 2")
    set_font(h2.add_run("1.1 Main Contributions"), size=12, bold=True)

    lead = doc.add_paragraph(style="Normal")
    set_font(lead.add_run("The main contributions of this study are as follows."), size=11)
    lead.paragraph_format.keep_with_next = True

    for idx, (label, detail) in enumerate(CONTRIBUTIONS, start=1):
        p = doc.add_paragraph(style="Contribution")
        set_font(p.add_run(f"{label} "), size=10.5, bold=True)
        set_font(p.add_run(detail), size=10.5)

    h2 = doc.add_paragraph(style="Heading 2")
    set_font(h2.add_run("1.2 Principal Findings and Implications"), size=12, bold=True)
    p = doc.add_paragraph(style="Normal")
    set_font(p.add_run(BODY_PARAGRAPHS[8]), size=11)

    if source_image:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as fh:
            fh.write(source_image)
            tmp_image = Path(fh.name)
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(15.4)
        cell = table.cell(0, 0)
        cell.width = Cm(15.4)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=80, start=0, bottom=50, end=0)
        pic_p = cell.paragraphs[0]
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.paragraph_format.keep_with_next = True
        pic_p.paragraph_format.space_after = Pt(2)
        pic_p.add_run().add_picture(str(tmp_image), width=Cm(15.2))
        remove_table_borders(table)
        tmp_image.unlink(missing_ok=True)

        cap = doc.add_paragraph(style="Caption")
        set_font(cap.add_run(FIGURE_CAPTION), size=9)

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = "Systematic Evaluation of Pluggable Preprocessing Pipelines for Heterogeneous WiFi CSI Sensing Datasets"
    core.subject = "Polished English translation of Chapter 1"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Translated and polished from the supplied Chinese draft; citation placeholders retained for author completion."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
